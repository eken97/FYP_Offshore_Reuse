"""
Build the standalone "Static Axial Capacity Check for Splash-Zone Members"
appendix as a .docx to paste into the thesis.

Everything is computed live from member_static_life_check.py /
sd_geometry.py against the real campaign CSVs, so the worked example and
the 32-row table are always self-consistent with the code. This version
uses the Design Basis GENERAL corrosion rate (0.30 mm/yr/surface) -- a
static/ULS check is not a fatigue check, so the halved 0.15 mm/yr fatigue
allowance does not apply here (fixed 28.08.2026).

Output: figures/static_check/Static_Member_Life_Check_Appendix.docx
"""
import math
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import member_static_life_check as mlc
import sd_geometry as sg

POSTPRO_DIR = Path(__file__).resolve().parent
PROJECT = POSTPRO_DIR.parent   # repo root
RESULTS_DIR = PROJECT / "results"
OUT_PATH = PROJECT / "figures" / "static_check" / "Static_Member_Life_Check_Appendix.docx"

RATES = sg.GENERAL_CORROSION_RATE_MM_PER_YEAR_PER_SURFACE
MATH_FONT = "Cambria Math"


# region --- intermediate-value calc (mirrors mlc.eurocode_capacity, but
#            returns every step for the worked example) ---
def intermediates(D_m, t_m, Lcr_m):
    D, t = D_m * 1000.0, t_m * 1000.0
    d = D - 2 * t
    Dt = D / t
    A = math.pi / 4 * (D**2 - d**2)
    cls = 1 if Dt <= 33 else 2 if Dt <= 46 else 3 if Dt <= 59 else 4
    A_eff = A * math.sqrt((90.0 / Dt) * (235.0 / mlc.FY_MPA)) if cls == 4 else A
    i_c = math.sqrt(D**2 + d**2) / 4.0
    Lcr = Lcr_m * 1000.0 * mlc.K_EFFECTIVE_LENGTH
    lam = Lcr / i_c
    lam_1 = math.pi * math.sqrt(mlc.E_MPA / mlc.FY_MPA)
    lam_bar = lam / lam_1
    psi = 0.5 * (1 + mlc.BUCKLING_CURVE_ALPHA * (lam_bar - 0.2) + lam_bar**2)
    chi = 1.0 / (psi + math.sqrt(psi**2 - lam_bar**2))
    N_b = chi * A_eff * mlc.FY_MPA / mlc.GAMMA_M
    N_c = A_eff * mlc.FY_MPA / mlc.GAMMA_M
    return dict(D=D, t=t, d=d, Dt=Dt, cls=cls, A=A, A_eff=A_eff, i_c=i_c,
               lam=lam, lam_1=lam_1, lam_bar=lam_bar, psi=psi, chi=chi,
               N_b=N_b, N_c=N_c)
# endregion


# region --- docx helpers ---
def _eq(doc, *parts):
    """Add an equation paragraph. Each part is a str (normal), ('sub', str)
    or ('sup', str)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for part in parts:
        if isinstance(part, tuple):
            kind, text = part
            r = p.add_run(text)
            r.font.subscript = (kind == "sub")
            r.font.superscript = (kind == "sup")
        else:
            r = p.add_run(part)
        r.font.name = MATH_FONT
        r.font.size = Pt(11)
    return p


def _num(x, nd=0):
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "--"
    if nd == 0:
        return f"{x:,.0f}"
    return f"{x:,.{nd}f}"
# endregion


def build():
    model = sg.read_subdyn_model(sg.DEFAULT_SD_PATH)
    forces = pd.read_csv(mlc.FORCE_CSV).set_index("member_id")
    static = pd.read_csv(mlc.OUT_CSV)

    doc = Document()
    doc.styles["Normal"].font.name = "Cambria"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("Static Axial Capacity Check for Splash-Zone Members", level=1)

    doc.add_paragraph(
        "This check bounds the fatigue-life extrapolation of the splash-zone "
        "members. As a member loses wall thickness to corrosion, its axial "
        "capacity falls; the static life reported here is the exposure time at "
        "which the corroded section first fails a tension or compression "
        "(buckling) check under its governing campaign load. The member "
        "stability check follows EN 1993-1-1 with the partial factors of "
        "DNV-ST-0126, as endorsed by that standard in place of DNV-RP-C202."
    )

    p = doc.add_paragraph()
    p.add_run("Corrosion rate. ").bold = True
    p.add_run(
        "This is an ultimate-limit-state capacity check, not a fatigue "
        "check, so the full UpWind Design Basis general rate of "
        "0.30 mm/yr per exposed surface is applied (0.30 mm/yr on both "
        "surfaces for the flooded legs; 0.30 mm/yr on the external surface "
        "only for the non-flooded braces). The halved 0.15 mm/yr allowance "
        "permitted for fatigue design is not used here."
    )

    # --- constants table ---
    doc.add_heading("Parameters", level=2)
    ct = doc.add_table(rows=1, cols=3)
    ct.style = "Table Grid"
    ct.rows[0].cells[0].text = "Symbol"
    ct.rows[0].cells[1].text = "Value"
    ct.rows[0].cells[2].text = "Source"
    for sym, val, src in [
        ("E", "210 000 MPa", "EN 1993-1-1 3.2.6"),
        ("f_y", "355 MPa (S355)", "EN 1993-1-1 / product standard"),
        ("gamma_M0 = gamma_M1", "1.10", "DNV-ST-0126 Table 4-8"),
        ("gamma_f", "1.35 (normal ULS)", "DNV-ST-0437"),
        ("alpha (buckling curve d)", "0.76", "EN 1993-1-1 Table 6.1; deliberate conservative choice for the splash-zone / reuse context"),
        ("K (effective-length factor)", "1.0", "pinned-pinned, member length between joints"),
        ("Corrosion rate (general)", "0.30 mm/yr/surface", "UpWind Design Basis"),
    ]:
        row = ct.add_row().cells
        row[0].text = sym
        row[1].text = val
        row[2].text = src

    # --- formula sheet ---
    doc.add_heading("Formulae", level=2)
    doc.add_paragraph(
        "For a circular hollow section of outer diameter D and wall "
        "thickness t (both after corrosion), critical length L_cr:"
    )
    _eq(doc, "(A.1)   d = D − 2t")
    _eq(doc, "(A.2)   section class:  D/t ≤ 33 → 1,   ≤ 46 → 2,   ≤ 59 → 3,   else 4")
    _eq(doc, "(A.3)   A = (π/4)(D² − d²)")
    _eq(doc, "(A.4)   A", ("sub", "eff"), " = A · √[ (90 / (D/t)) · (235 / f", ("sub", "y"), ") ]      (Class 4 only; otherwise A", ("sub", "eff"), " = A)")
    _eq(doc, "(A.5)   i", ("sub", "c"), " = √(D² + d²) / 4")
    _eq(doc, "(A.6)   L", ("sub", "cr"), " = K · L ,   K = 1.0")
    _eq(doc, "(A.7)   λ = L", ("sub", "cr"), " / i", ("sub", "c"))
    _eq(doc, "(A.8)   λ", ("sub", "1"), " = π · √(E / f", ("sub", "y"), ")")
    _eq(doc, "(A.9)   λ̄ = λ / λ", ("sub", "1"), "      (relative slenderness)")
    _eq(doc, "(A.10)  ψ = 0.5 [ 1 + α(λ̄ − 0.2) + λ̄² ] ,   α = 0.76 (curve d)")
    _eq(doc, "(A.11)  χ = 1 / ( ψ + √(ψ² − λ̄²) )")
    _eq(doc, "(A.12)  N", ("sub", "b,Rd"), " = χ · A", ("sub", "eff"), " · f", ("sub", "y"), " / γ", ("sub", "M1"), "      (compression / buckling)")
    _eq(doc, "(A.13)  N", ("sub", "c,Rd"), " = A", ("sub", "eff"), " · f", ("sub", "y"), " / γ", ("sub", "M1"), "      (tension)")
    _eq(doc, "(A.14)  N", ("sub", "Ed"), " = γ", ("sub", "f"), " · F      — the section fails when N", ("sub", "Ed,c"), " > N", ("sub", "b,Rd"), "  or  N", ("sub", "Ed,t"), " > N", ("sub", "c,Rd"))

    doc.add_paragraph(
        "Corrosion is stepped forward in time (0.1-year resolution). After "
        "τ years of exposure at rate r per surface:"
    )
    _eq(doc, "(A.15)  D(τ) = D", ("sub", "0"), " − 2 r", ("sub", "ext"), " τ / 1000")
    _eq(doc, "(A.16)  t(τ) = t", ("sub", "0"), " − (r", ("sub", "ext"), " + r", ("sub", "int"), ") τ / 1000")
    doc.add_paragraph(
        "with r_ext = r_int = 0.30 mm/yr for legs and r_ext = 0.30 mm/yr, "
        "r_int = 0 for braces. The static life is the smallest τ for which "
        "(A.14) is violated; the governing mode is whichever of the two "
        "utilisations reaches 1.0 first."
    )

    # --- worked example: member 18 ---
    doc.add_heading("Worked example — Member 18 (leg)", level=2)
    mid = 18
    m = model["members"][mid]
    p1 = model["joints"][m["j1"]]
    p2 = model["joints"][m["j2"]]
    L_m = math.dist(p1, p2)
    Fc = forces.loc[mid, "max_compressive_N"]
    Ft = forces.loc[mid, "max_tensile_N"]
    row = static.set_index("member_id").loc[mid]
    yr_fail = row["static_life_years"]

    doc.add_paragraph(
        f"Member 18 is a splash-zone leg, D0 = 1200 mm, t0 = 35 mm, "
        f"length between joints L = {L_m:.2f} m. Its governing campaign load "
        f"is compression, F = {abs(Fc):,.0f} N, giving a design demand "
        f"N_Ed,c = gamma_f · F = 1.35 × {abs(Fc):,.0f} = "
        f"{mlc.GAMMA_F * abs(Fc):,.0f} N. "
        f"The table compares the section as-new against the section at the "
        f"failure year, τ = {yr_fail:.1f} yr."
    )

    d0 = sg.corroded_section(model, mid, 0.0, rates=RATES)
    df = sg.corroded_section(model, mid, yr_fail, rates=RATES)
    i0 = intermediates(*d0, L_m)
    i1 = intermediates(*df, L_m)

    wt = doc.add_table(rows=1, cols=3)
    wt.style = "Table Grid"
    hdr = wt.rows[0].cells
    hdr[0].text = "Quantity"
    hdr[1].text = "τ = 0 yr"
    hdr[2].text = f"τ = {yr_fail:.1f} yr"
    spec = [
        ("D  (mm)", "D", 2), ("t  (mm)", "t", 2), ("d  (mm)", "d", 2),
        ("D / t", "Dt", 2), ("section class", "cls", 0),
        ("A  (mm²)", "A", 0), ("A_eff  (mm²)", "A_eff", 0),
        ("i_c  (mm)", "i_c", 1), ("λ", "lam", 2),
        ("λ̄", "lam_bar", 3), ("ψ", "psi", 3), ("χ", "chi", 3),
        ("N_b,Rd  (N)", "N_b", 0), ("N_c,Rd  (N)", "N_c", 0),
    ]
    for label, key, nd in spec:
        c = wt.add_row().cells
        c[0].text = label
        c[1].text = _num(i0[key], nd) if key != "cls" else str(i0[key])
        c[2].text = _num(i1[key], nd) if key != "cls" else str(i1[key])
    c = wt.add_row().cells
    c[0].text = "N_Ed,c  (N)"
    c[1].text = f"{mlc.GAMMA_F * abs(Fc):,.0f}"
    c[2].text = f"{mlc.GAMMA_F * abs(Fc):,.0f}"

    doc.add_paragraph(
        f"As-new, N_b,Rd = {i0['N_b']:,.0f} N far exceeds the "
        f"{mlc.GAMMA_F * abs(Fc):,.0f} N demand. By τ = {yr_fail:.1f} yr the "
        f"section has lost enough wall (t down to {i1['t']:.1f} mm, now "
        f"Class 4) that N_b,Rd = {i1['N_b']:,.0f} N has fallen just below the "
        f"demand — the compression check fails and the static life is "
        f"{yr_fail:.1f} yr. Note that t at failure, and every quantity "
        f"derived from it, is fixed by load and geometry alone; only the "
        f"time taken to corrode to that thickness depends on the assumed "
        f"rate."
    )

    # --- full results table ---
    doc.add_heading("Static life results — all 32 splash-zone members", level=2)
    doc.add_paragraph(
        "Governing force is the campaign extreme (tension positive, "
        "compression negative) driving the failing check. t_fail / D_fail "
        "are the section dimensions at the failure year."
    )
    rt = doc.add_table(rows=1, cols=8)
    rt.style = "Table Grid"
    for i, h in enumerate([
        "Member", "Class", "L (m)", "Governing force (N)",
        "Mode", "t_fail (mm)", "D_fail (mm)", "Static life (yr)",
    ]):
        rt.rows[0].cells[i].text = h

    # legs first (they carry the lower, governing lives), then by member id
    # -- an explicit key, not string order, which would put "brace" first
    static["_cls_order"] = static["member_class"].map({"leg": 0, "brace": 1})
    sdf = static.sort_values(["_cls_order", "member_id"])
    for _, r in sdf.iterrows():
        gov_force = r["max_compressive_N"] if r["governing_mode"] == "compression" else r["max_tensile_N"]
        c = rt.add_row().cells
        c[0].text = str(int(r["member_id"]))
        c[1].text = r["member_class"]
        c[2].text = f"{r['L_m']:.2f}"
        c[3].text = f"{gov_force:,.0f}"
        c[4].text = r["governing_mode"]
        c[5].text = f"{r['t_at_failure_mm']:.2f}"
        c[6].text = f"{r['D_at_failure_mm']:.1f}"
        c[7].text = f"{r['static_life_years']:.1f}"

    legs = sdf[sdf.member_class == "leg"]["static_life_years"]
    braces = sdf[sdf.member_class == "brace"]["static_life_years"]
    doc.add_paragraph(
        f"Range: legs {legs.min():.1f}–{legs.max():.1f} yr "
        f"(median {legs.median():.1f}), braces {braces.min():.1f}–"
        f"{braces.max():.1f} yr (median {braces.median():.1f}). "
        f"The smallest static life is {sdf['static_life_years'].min():.1f} yr "
        f"(Member {int(sdf.loc[sdf['static_life_years'].idxmin(), 'member_id'])}), "
        f"still clearing the 25-year design life by a factor of about "
        f"{sdf['static_life_years'].min() / 25:.1f}. "
        f"The static check governs for all 32 of 32 splash-zone members: "
        f"the fatigue-extrapolated lives (160–680 yr) are never reached."
    )

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            continue
        para.paragraph_format.space_after = Pt(6)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")
    print(f"  worked example: member 18 fails compression at {yr_fail:.1f} yr")
    print(f"  results table: {len(sdf)} rows, "
          f"min life {sdf['static_life_years'].min():.1f} yr")


if __name__ == "__main__":
    build()
